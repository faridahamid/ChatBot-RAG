import os
import uuid
import hashlib
import re
import io
from typing import List, Optional
import pandas as pd
from pypdf import PdfReader
import docx
from sqlalchemy.orm import Session
from sentence_transformers import SentenceTransformer
from models import Document, DocumentChunk

# from FlagEmbedding import BGEM3FlagModel
import numpy as np

EMBEDDING_MODEL = os.getenv(
    "EMBEDDING_MODEL",
    "sentence-transformers/distiluse-base-multilingual-cased-v2"
)

# Set to "cuda" to use GPU if available, otherwise "cpu"
ST_DEVICE = os.getenv("EMBEDDING_DEVICE", "cpu").strip().lower()  # "cpu" | "cuda"
ST_BATCH_SIZE = int(os.getenv("EMBEDDING_BATCH_SIZE", "64"))      # effective encode batch

_model: Optional[SentenceTransformer] = None
try:
    print(f"Loading embedding model: {EMBEDDING_MODEL} on device={ST_DEVICE}")
    _model = SentenceTransformer(EMBEDDING_MODEL, device=ST_DEVICE)
    print("Embedding model loaded successfully")
except Exception as e:
    raise RuntimeError(f"Could not load embedding model: {e}")

# ---------- File readers ----------
def _read_pdf(path: str) -> str:
    reader = PdfReader(path)
    return "\n".join([(p.extract_text() or "") for p in reader.pages])

def _read_docx(path: str) -> str:
    d = docx.Document(path)
    return "\n".join([para.text for para in d.paragraphs])

def _read_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def _read_csv(path: str) -> str:
    df = pd.read_csv(path)
    df = df.fillna("")
    lines = []
    for _, row in df.iterrows():
        line = " | ".join(f"{col}: {row[col]}" for col in df.columns)
        lines.append(line)
    return "\n".join(lines)

# In-memory readers
def _read_pdf_bytes(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    return "\n".join([(p.extract_text() or "") for p in reader.pages])

def _read_docx_bytes(data: bytes) -> str:
    d = docx.Document(io.BytesIO(data))
    return "\n".join([para.text for para in d.paragraphs])

def _read_txt_bytes(data: bytes) -> str:
    return data.decode("utf-8", errors="ignore")

def _read_csv_bytes(data: bytes) -> str:
    df = pd.read_csv(io.BytesIO(data))
    df = df.fillna("")
    lines = []
    for _, row in df.iterrows():
        line = " | ".join(f"{col}: {row[col]}" for col in df.columns)
        lines.append(line)
    return "\n".join(lines)

def extract_text(file_path: str, filetype: str) -> str:
    ft = (filetype or "").lower()
    if ft == "pdf":
        return _read_pdf(file_path)
    if ft == "docx":
        return _read_docx(file_path)
    if ft == "txt":
        return _read_txt(file_path)
    if ft == "csv":
        return _read_csv(file_path)
    raise ValueError(f"Unsupported file type: {filetype}")

def extract_text_from_bytes(file_bytes: bytes, filetype: str) -> str:
    ft = (filetype or "").lower()
    print(f"Extracting text from {ft} file ({len(file_bytes)} bytes)")
    
    try:
        if ft == "pdf":
            text = _read_pdf_bytes(file_bytes)
        elif ft == "docx":
            text = _read_docx_bytes(file_bytes)
        elif ft == "txt":
            text = _read_txt_bytes(file_bytes)
        elif ft == "csv":
            text = _read_csv_bytes(file_bytes)
        else:
            raise ValueError(f"Unsupported file type: {filetype}")
        
        print(f"Extracted {len(text)} characters of text")
        if len(text.strip()) == 0:
            print("Warning: Extracted text is empty")
        return text
    except Exception as e:
        print(f"Error extracting text from {ft} file: {e}")
        raise Exception(f"Failed to extract text from {ft} file: {e}")

# ---------- Chunk & Embed ----------
def chunk_text(text: str, max_chars: int = 800, overlap: int = 100) -> List[str]:
    text = (text or "").strip()
    print(f"Chunking text of length {len(text)} with max_chars={max_chars}, overlap={overlap}")
    
    if not text:
        print("Warning: Empty text provided for chunking")
        return []
    
    chunks, i, n = [], 0, len(text)
    step = max(1, max_chars - overlap)
    
    while i < n:
        end = min(i + max_chars, n)
        piece = text[i:end].strip()
        if piece:
            chunks.append(piece)
        i += step
    
    print(f"Created {len(chunks)} chunks")
    return chunks

# def _embed_passages_batch(texts: List[str]) -> List[List[float]]:
    
#     try:
#         print(f"Embedding {len(texts)} text chunks...")
#         out = _model.encode(
#             texts,
#             batch_size=max(1, len(texts)),
#             return_dense=True,
#             return_sparse=False,
#             return_colbert_vecs=False,
#         )
#         vecs = out["dense_vecs"]
#         # Convert to numpy array and L2-normalize
#         vecs_np = np.array(vecs, dtype=np.float32)
#         norms = np.linalg.norm(vecs_np, axis=1, keepdims=True)
#         norms = np.clip(norms, 1e-12, None)
#         vecs_np = vecs_np / norms
#         result = vecs_np.tolist()
#         print(f"Successfully embedded {len(result)} chunks")
#         return result
#     except Exception as e:
#         print(f"Error during embedding: {e}")
#         raise Exception(f"Embedding failed: {e}")

def _embed_passages_batch(texts: List[str]) -> List[List[float]]:
    try:
        print(f"Embedding {len(texts)} text chunks...")
        vecs = _model.encode(
            texts,
            batch_size=max(1, len(texts)),
            normalize_embeddings=True,   # auto L2 normalize
        )
        result = vecs.tolist()
        print(f"Successfully embedded {len(result)} chunks")
        return result
    except Exception as e:
        print(f"Error during embedding: {e}")
        raise Exception(f"Embedding failed: {e}")

# def embed_query(question: str) -> List[float]:
 
#     out = _model.encode(
#         [question],
#         batch_size=1,
#         return_dense=True,
#         return_sparse=False,
#         return_colbert_vecs=False,
#     )
#     v = out["dense_vecs"][0]
#     v_np = np.array(v, dtype=np.float32)
#     norm = float(np.linalg.norm(v_np))
#     norm = max(norm, 1e-12)
#     v_np = v_np / norm
#     return v_np.tolist()

def embed_query(question: str) -> List[float]:
    v = _model.encode(
        [question],
        batch_size=1,
        normalize_embeddings=True,
    )[0]
    return v.tolist()

# ---------- Hash helpers ----------
def extract_text_from_raw(raw_text: Optional[str]) -> str:
    return (raw_text or "").strip()

def _normalize_text_for_hash(text: str) -> str:
    lowered = text.lower()
    collapsed = re.sub(r"\s+", " ", lowered)
    return collapsed.strip()

def _sha256_hex(content: str) -> str:
    return hashlib.sha256(content.encode("utf-8", errors="ignore")).hexdigest()

# ---------- Ingestion pipelines ----------
def process_document(
    db: Session,
    org_id: str,
    user_id: str,
    file_path: str,
    filetype: str,
    chunk_size: int = 800,
    chunk_overlap: int = 100,
    embedding_batch_size: int = 64,
    insert_batch_size: int = 200,
):
    # Read + dedupe per org
    text_all = extract_text(file_path, filetype)
    normalized = _normalize_text_for_hash(text_all)
    content_hash = _sha256_hex(normalized)

    existing = (
        db.query(Document)
        .filter(Document.organization_id == org_id, Document.content_hash == content_hash)
        .first()
    )
    if existing is not None:
        raise ValueError("DUPLICATE_DOCUMENT")

    doc = Document(
        id=uuid.uuid4(),
        organization_id=org_id,
        uploaded_by=user_id,
        filename=os.path.basename(file_path),
        filetype=filetype.lower(),
        content_hash=content_hash,
    )
    db.add(doc)
    # Flush to obtain doc.id without committing the transaction yet
    db.flush()

    chunks = chunk_text(text_all, max_chars=chunk_size, overlap=chunk_overlap)
    total = len(chunks)
    
    print(f"Created {total} chunks from document")

    if total == 0:
        print("Warning: No chunks created from document")
        return {"document_id": doc.id, "chunks": 0}

    try:
        for start in range(0, total, embedding_batch_size):
            end = min(start + embedding_batch_size, total)
            batch_texts = chunks[start:end]
            print(f"Processing batch {start//embedding_batch_size + 1}: chunks {start+1}-{end}")
            
            batch_vecs = _embed_passages_batch(batch_texts)

            batch_objects: List[DocumentChunk] = [
                DocumentChunk(
                    id=uuid.uuid4(),
                    document_id=doc.id,
                    content=content,
                    embedding=vec,
                )
                for content, vec in zip(batch_texts, batch_vecs)
            ]

            for i in range(0, len(batch_objects), insert_batch_size):
                slice_objs = batch_objects[i : i + insert_batch_size]
                if not slice_objs:
                    continue
                db.bulk_save_objects(slice_objs)
                db.flush()
                db.commit()
                print(f"Saved {len(slice_objs)} chunks to database")
    except Exception as e:
        print(f"Error during chunk processing: {e}")
        # Rollback the document creation if chunking fails
        db.rollback()
        raise Exception(f"Failed to process document chunks: {e}")

    return {"document_id": doc.id, "chunks": total}

def process_document_from_bytes(
    db: Session,
    org_id: str,
    user_id: str,
    file_bytes: bytes,
    filename: str,
    filetype: str,
    chunk_size: int = 800,
    chunk_overlap: int = 100,
    embedding_batch_size: int = 64,
    insert_batch_size: int = 200,
):
    text_all = extract_text_from_bytes(file_bytes, filetype)
    normalized = _normalize_text_for_hash(text_all)
    content_hash = _sha256_hex(normalized)

    existing = (
        db.query(Document)
        .filter(Document.organization_id == org_id, Document.content_hash == content_hash)
        .first()
    )
    if existing is not None:
        raise ValueError("DUPLICATE_DOCUMENT")

    doc = Document(
        id=uuid.uuid4(),
        organization_id=org_id,
        uploaded_by=user_id,
        filename=filename,
        filetype=filetype.lower(),
        content_hash=content_hash,
    )
    db.add(doc)
    db.flush()

    chunks = chunk_text(text_all, max_chars=chunk_size, overlap=chunk_overlap)
    total = len(chunks)
    
    print(f"Created {total} chunks from document")

    if total == 0:
        print("Warning: No chunks created from document")
        return {"document_id": doc.id, "chunks": 0}

    try:
        for start in range(0, total, embedding_batch_size):
            end = min(start + embedding_batch_size, total)
            batch_texts = chunks[start:end]
            print(f"Processing batch {start//embedding_batch_size + 1}: chunks {start+1}-{end}")
            
            batch_vecs = _embed_passages_batch(batch_texts)

            batch_objects: List[DocumentChunk] = [
                DocumentChunk(
                    id=uuid.uuid4(),
                    document_id=doc.id,
                    content=content,
                    embedding=vec,
                )
                for content, vec in zip(batch_texts, batch_vecs)
            ]
            for i in range(0, len(batch_objects), insert_batch_size):
                slice_objs = batch_objects[i : i + insert_batch_size]
                if not slice_objs:
                    continue
                db.bulk_save_objects(slice_objs)
                db.flush()
                db.commit()
                print(f"Saved {len(slice_objs)} chunks to database")
    except Exception as e:
        print(f"Error during chunk processing: {e}")
        # Rollback the document creation if chunking fails
        db.rollback()
        raise Exception(f"Failed to process document chunks: {e}")

    return {"document_id": doc.id, "chunks": total}
